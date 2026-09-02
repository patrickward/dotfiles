#!/usr/bin/env python3
"""
Generate a pandoc-compatible reference.docx with custom styles.
Usage:
  python3 gen_reference_docx.py              # Black/gray theme
  python3 gen_reference_docx.py --blue       # Blue theme
"""

import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- COLOR SCHEMES (only styles python-docx natively supports) ---

COLORS_BLACK = {
    "Normal": RGBColor(0, 0, 0),
    "Body Text": RGBColor(0, 0, 0),
    "Heading 1": RGBColor(20, 20, 20),
    "Heading 2": RGBColor(40, 40, 40),
    "Heading 3": RGBColor(60, 60, 60),
    "Heading 4": RGBColor(70, 70, 70),
    "Heading 5": RGBColor(80, 80, 80),
    "Heading 6": RGBColor(90, 90, 90),
    "Title": RGBColor(30, 30, 30),
    "Subtitle": RGBColor(50, 50, 50),
    "Caption": RGBColor(60, 60, 60),
}

COLORS_BLUE = {
    "Normal": RGBColor(0, 0, 0),
    "Body Text": RGBColor(0, 0, 0),
    "Heading 1": RGBColor(0, 30, 60),
    "Heading 2": RGBColor(0, 50, 100),
    "Heading 3": RGBColor(0, 70, 140),
    "Heading 4": RGBColor(0, 90, 180),
    "Heading 5": RGBColor(0, 110, 200),
    "Heading 6": RGBColor(0, 130, 220),
    "Title": RGBColor(0, 40, 80),
    "Subtitle": RGBColor(0, 60, 120),
    "Caption": RGBColor(0, 70, 140),
}

# Styles that python-docx natively supports and we want to customize
STYLES_TO_MODIFY = [
    ("Normal", True),
    ("Body Text", True),
    ("Heading 1", True),
    ("Heading 2", True),
    ("Heading 3", True),
    ("Heading 4", True),
    ("Heading 5", True),
    ("Heading 6", True),
    ("Title", True),
    ("Subtitle", True),
    ("Caption", True),
]

def apply_styles(doc, colors):
    """Apply custom colors and left-alignment to supported styles."""
    for style_name, is_para in STYLES_TO_MODIFY:
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        if is_para:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style_name in colors:
            style.font.color.rgb = colors[style_name]

def add_sample_content(doc):
    """Add sample content for all supported styles."""
    doc.add_heading("Title Style", level=0)

    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)
    doc.add_heading("Heading 4", level=4)
    doc.add_heading("Heading 5", level=5)
    doc.add_heading("Heading 6", level=6)

    doc.add_paragraph("Normal paragraph text.")

    doc.add_paragraph("Caption text", style="Caption")

    para = doc.add_paragraph()
    para.add_run("Normal text with a ")
    hyperlink = para.add_run("hyperlink")

def main():
    use_blue = "--blue" in sys.argv
    colors = COLORS_BLUE if use_blue else COLORS_BLACK
    theme_name = "blue" if use_blue else "black"

    doc = Document()
    apply_styles(doc, colors)
    add_sample_content(doc)

    filename = f"pandoc-reference-{theme_name}.docx"
    doc.save(filename)
    print(f"Created: {filename}")

if __name__ == "__main__":
    main()